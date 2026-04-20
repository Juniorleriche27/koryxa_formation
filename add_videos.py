import json
import urllib.parse
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def youtube_search_url(titre):
    return "https://www.youtube.com/results?search_query=" + urllib.parse.quote(titre)

# ── Données vidéos ────────────────────────────────────────────────────────────
videos = {
    "MODULE_1_Bases_Python_Data.ipynb": [
        (
            "Maîtrisez les bases de Python en 50 minutes — 15 exercices pour débutants",
            "Variables, types, fonctions, boucles mis en pratique à travers 15 exercices progressifs. "
            "Idéal pour consolider tout ce que tu viens d'apprendre dans ce module.",
        ),
        (
            "Apprenez Python facilement : cours complet pour débutants",
            "Conditions, boucles et fonctions expliqués pour des projets réels. "
            "Le formateur explique chaque concept avec des exemples du quotidien — parfait pour les débutants.",
        ),
    ],
    "MODULE_2_NumPy_Calcul_Numerique.ipynb": [
        (
            "[Tuto Français] Introduction à Python NumPy",
            "Arrays, opérations vectorisées, indexation et slicing expliqués étape par étape. "
            "Le formateur code en direct — tu peux suivre en même temps dans ton propre notebook.",
        ),
        (
            "Débuter avec la librairie NumPy sur Python",
            "Prise en main progressive avec les opérations essentielles. "
            "Parfait pour voir NumPy en action sur des exemples différents de ceux de ce module.",
        ),
    ],
    "MODULE_3_Pandas_Manipulation_Donnees.ipynb": [
        (
            "Maîtriser Pandas Pour La Data Science En 20 min — Guide Complet Débutants",
            "DataFrames, filtres et groupby expliqués en 20 minutes chrono. "
            "Excellent résumé visuel de tout ce que tu viens d'apprendre — idéal après avoir fini le module.",
        ),
        (
            "PANDAS PYTHON Français — Introduction + Analyse du Titanic",
            "Pandas mis en action sur le célèbre dataset Titanic. "
            "Voir Pandas appliqué à un vrai dataset différent du nôtre renforce beaucoup la compréhension.",
        ),
    ],
    "MODULE_4_Nettoyage_Donnees.ipynb": [
        (
            "Data Cleaning avec Python et Pandas : Le Guide Complet",
            "1 heure de cours entièrement dédié au nettoyage de données de A à Z. "
            "Ce tutoriel va plus loin que notre module et couvre des cas avancés que tu rencontreras en entreprise.",
        ),
        (
            "Nettoyer et analyser votre jeu de données avec Python — Épisode 1 : Introduction",
            "Série pratique sur les valeurs manquantes, erreurs et types. "
            "Le format en épisodes te permet de regarder à ton rythme et de cibler les points qui te posent problème.",
        ),
    ],
    "MODULE_5_Visualisation_Donnees.ipynb": [
        (
            "Initiation à Matplotlib pour la visualisation des données avec Python",
            "Courbes, barres, histogrammes et personnalisation des graphiques. "
            "Voir le formateur construire les graphiques en direct est beaucoup plus clair que de lire la documentation.",
        ),
        (
            "Tutoriel Seaborn pour Débutants — Les plus beaux graphiques en une ligne de code",
            "Heatmap, boxplot, pairplot avec Seaborn — des graphiques professionnels en très peu de code. "
            "La vidéo montre le résultat visuel immédiatement, ce qui aide à choisir le bon graphique.",
        ),
    ],
    "MODULE_6_Analyse_Exploratoire_EDA.ipynb": [
        (
            "Réaliser une analyse exploratoire avec Python — Pandas, pandas-profiling, Plotly",
            "EDA complète sur un dataset réel avec des outils modernes. "
            "Tu découvriras pandas-profiling qui génère automatiquement un rapport EDA complet — très utile pour le Module 7.",
        ),
        (
            "Guide de l'Analyse Exploratoire des Données (EDA) étape par étape",
            "Méthodologie structurée du début à la fin d'une EDA. "
            "Ce tutoriel montre comment poser les bonnes questions avant de coder — c'est la compétence clé d'un data analyst.",
        ),
    ],
}

videos_module0 = [
    (
        "Installation et prise en main d'Anaconda / Jupyter Notebook / Python",
        "Installation complète + premier code Python dans Jupyter. "
        "Si tu bloques lors de l'installation, cette vidéo te montre chaque étape en détail — "
        "tu peux la mettre en pause et reproduire exactement ce que fait le formateur.",
    ),
    (
        "Les bases de la data science avec Python",
        "Vue d'ensemble de l'environnement et de l'écosystème data Python. "
        "Idéal pour comprendre la big picture avant de commencer les modules — "
        "tu sauras exactement où tu vas avant de commencer.",
    ),
]

# ── AJOUTER DANS LES NOTEBOOKS ────────────────────────────────────────────────
for fichier, vids in videos.items():
    with open(f"c:/KORYXA_formation/{fichier}", encoding="utf-8") as f:
        nb = json.load(f)

    # Supprimer ancienne cellule vidéos si elle existe
    nb["cells"] = [c for c in nb["cells"] if c.get("id") != "videos_youtube"]

    lignes = [
        "---\n",
        "---\n",
        "# 🎬 Vidéos YouTube — Pour aller plus loin\n\n",
        "Ces 2 vidéos en français complètent parfaitement ce module. "
        "Regarde-les **après** avoir terminé le notebook pour consolider ce que tu as appris.\n\n",
        "---\n\n",
    ]
    for i, (titre, description) in enumerate(vids, 1):
        url = youtube_search_url(titre)
        lignes += [
            f"### 🎥 Vidéo {i} — {titre}\n\n",
            f"**Rechercher sur YouTube :** [{titre}]({url})\n\n",
            f"{description}\n\n",
            "---\n\n",
        ]
    lignes.append(
        "> 💡 **Conseil :** Regarde d'abord le notebook en entier, fais les exercices, "
        "**puis** regarde les vidéos. Les vidéos servent à consolider, pas à remplacer la pratique."
    )

    nb["cells"].append({
        "cell_type": "markdown",
        "id": "videos_youtube",
        "metadata": {},
        "source": lignes,
    })

    with open(f"c:/KORYXA_formation/{fichier}", "w", encoding="utf-8") as f:
        json.dump(nb, f, ensure_ascii=False, indent=1)

    print(f"Videos ajoutees : {fichier}")

# ── AJOUTER DANS LE DOCUMENT WORD ────────────────────────────────────────────
def set_color(run, r, g, b):
    run.font.color.rgb = RGBColor(r, g, b)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2196F3")
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document("c:/KORYXA_formation/MODULE_0_Introduction_Installation.docx")

# Page break + grand titre
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run("Vidéos YouTube recommandées")
run.bold = True
run.font.size = Pt(18)
set_color(run, 33, 37, 41)
p.paragraph_format.space_after = Pt(6)
add_hr(doc)

p = doc.add_paragraph()
run = p.add_run(
    "14 vidéos en français, 2 par module, pour voir chaque concept expliqué en direct par un formateur. "
    "Ces vidéos complètent les notebooks — regardez-les après avoir terminé chaque module."
)
run.font.size = Pt(11)
run.italic = True
p.paragraph_format.space_after = Pt(10)

# Tous les modules avec leurs vidéos
tous_modules = [
    ("MODULE 0 — Introduction & Installation", videos_module0),
    ("MODULE 1 — Les Bases de Python pour la Data",   videos["MODULE_1_Bases_Python_Data.ipynb"]),
    ("MODULE 2 — NumPy : Calcul Numérique",           videos["MODULE_2_NumPy_Calcul_Numerique.ipynb"]),
    ("MODULE 3 — Pandas : Manipulation de Données",   videos["MODULE_3_Pandas_Manipulation_Donnees.ipynb"]),
    ("MODULE 4 — Nettoyage de Données",               videos["MODULE_4_Nettoyage_Donnees.ipynb"]),
    ("MODULE 5 — Visualisation de Données",           videos["MODULE_5_Visualisation_Donnees.ipynb"]),
    ("MODULE 6 — Analyse Exploratoire (EDA)",         videos["MODULE_6_Analyse_Exploratoire_EDA.ipynb"]),
]

for module_titre, vids in tous_modules:
    # Titre module
    p = doc.add_paragraph()
    run = p.add_run(module_titre)
    run.bold = True
    run.font.size = Pt(14)
    set_color(run, 33, 150, 243)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

    for i, (titre, description) in enumerate(vids, 1):
        url = youtube_search_url(titre)

        # Numéro + titre vidéo
        p = doc.add_paragraph()
        r1 = p.add_run(f"  Vidéo {i} — ")
        r1.bold = True
        r1.font.size = Pt(11)
        set_color(r1, 76, 175, 80)
        r2 = p.add_run(titre)
        r2.bold = True
        r2.font.size = Pt(11)
        set_color(r2, 33, 37, 41)
        p.paragraph_format.space_after = Pt(2)

        # Lien YouTube
        p = doc.add_paragraph()
        r1 = p.add_run("  Rechercher sur YouTube : ")
        r1.font.size = Pt(10)
        r1.bold = True
        set_color(r1, 100, 100, 100)
        r2 = p.add_run(url)
        r2.font.size = Pt(10)
        set_color(r2, 255, 0, 0)  # rouge YouTube
        r2.underline = True
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)

        # Description
        p = doc.add_paragraph()
        run = p.add_run(f"  {description}")
        run.font.size = Pt(10)
        run.italic = True
        set_color(run, 80, 80, 80)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)

    add_hr(doc)

# Conseil final
p = doc.add_paragraph()
run = p.add_run(
    "Conseil : Regarde chaque vidéo APRÈS avoir terminé le notebook correspondant. "
    "Les vidéos servent à consolider ce que tu as appris, pas à le remplacer. "
    "La pratique sur les notebooks est irremplaçable."
)
run.bold = True
run.italic = True
run.font.size = Pt(11)
set_color(run, 255, 152, 0)

doc.save("c:/KORYXA_formation/MODULE_0_Introduction_Installation.docx")
print("Document Word mis a jour avec les videos.")
print("Termine !")
